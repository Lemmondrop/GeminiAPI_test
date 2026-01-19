import os
import glob
import json
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from processor import refine_pdf_to_json_onecall, JSON_SCHEMA

# =========================================================
# ✅ [설정]
# =========================================================
plt.rcParams["font.family"] = "Malgun Gothic"
plt.rcParams["axes.unicode_minus"] = False

COLOR_RED = "indianred"
COLOR_YELLOW = "khaki"
COLOR_BLUE = "cornflowerblue"

# =========================================================
# 1. 문서 스타일링 헬퍼
# =========================================================
def set_cell_background(cell, fill_color):
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def apply_table_style(table, header_color="E7E6E6", center_all=True):
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if center_all:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i == 0: 
                set_cell_background(cell, header_color)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

# =========================================================
# 2. 데이터 처리 및 단위 감지
# =========================================================
def clean_numeric(val):
    if isinstance(val, (int, float)):
        return float(val)
    if val is None:
        return 0.0
    s = str(val).strip()
    if s.startswith("(") and s.endswith(")"):
        s = "-" + s[1:-1]
    s = s.replace(",", "").replace("억", "").replace("원", "").replace("개", "").replace("만", "").replace("불", "").replace("$", "")
    if s in ["", "N/A", "-", "n/a", "null"]:
        return 0.0
    try:
        return float(s)
    except:
        return 0.0

def detect_unit_from_data(data_list, default_unit="(단위: 백만원)"):
    if not data_list: return default_unit
    if isinstance(data_list[0], list):
        flat_str = " ".join([str(row[1]) for row in data_list[1:] if len(row) > 1])
    else:
        flat_str = " ".join([str(v) for v in data_list])
    if "억" in flat_str: return "(단위: 억원)"
    if "조" in flat_str: return "(단위: 조원)"
    if "천만" in flat_str: return "(단위: 천만원)"
    if "백만" in flat_str: return "(단위: 백만원)"
    if "$" in flat_str or "달러" in flat_str: return "(단위: USD)"
    if "개" in flat_str or "건" in flat_str: return "(단위: 건/개)"
    return default_unit

# =========================================================
# 3. 그래프 생성 함수
# =========================================================
def create_basic_bar_chart(data_list, title, color, file_name, suffix, target_dir, default_unit_label="(단위: 자료 수치 기준)"):
    if not data_list or len(data_list) < 2: return None
    unit_label = detect_unit_from_data(data_list, default_unit_label)

    try:
        years, values = [], []
        for row in data_list[1:]:
            if isinstance(row, list) and len(row) >= 2:
                y = str(row[0]).replace("년","").strip()
                v = clean_numeric(row[1])
                years.append(y)
                values.append(v)
        
        if not years: return None

        plt.figure(figsize=(7, 4.5))
        if len(years) == 1:
            plt.bar(years, values, color=color, width=0.3)
            plt.xlim(-1, 1)
        else:
            plt.bar(years, values, color=color, width=0.5)

        plt.title(f"{title}\n", fontsize=14, fontweight="bold", pad=10)
        plt.text(0, 1.02, unit_label, transform=plt.gca().transAxes, ha='left', fontsize=10, color='gray')
        plt.gca().yaxis.set_major_formatter(ticker.FuncFormatter(lambda x, p: format(int(x), ',')))
        plt.grid(axis="y", linestyle="--", alpha=0.5)
        
        max_val = max(values) if values else 1
        for bar in plt.gca().patches:
            height = bar.get_height()
            if height != 0:
                plt.text(bar.get_x() + bar.get_width()/2.0, height + (max_val*0.01),
                        f"{int(height):,}", ha='center', va='bottom', fontsize=9, fontweight='bold')

        save_path = os.path.join(target_dir, f"{file_name}_{suffix}.png")
        plt.tight_layout()
        plt.savefig(save_path, dpi=150)
        plt.close()
        return save_path
    except Exception as e:
        print(f"  [Graph Error] {title}: {e}")
        return None

def create_grouped_bar_chart_3(years, data1, data2, data3, labels, colors, title, file_name, suffix, target_dir, unit_label="(단위: 백만원)"):
    try:
        if not years or len(years) == 0: return None
        target_len = len(years)
        def pad(d):
            if d is None: d = []
            if len(d) < target_len: return d + [0]*(target_len - len(d))
            return d[:target_len]

        d1 = [clean_numeric(v) for v in pad(data1)]
        d2 = [clean_numeric(v) for v in pad(data2)]
        d3 = [clean_numeric(v) for v in pad(data3)]

        x = np.arange(len(years))
        width = 0.25

        plt.figure(figsize=(8.5, 5))
        plt.bar(x - width, d1, width, label=labels[0], color=colors[0])
        plt.bar(x,         d2, width, label=labels[1], color=colors[1])
        plt.bar(x + width, d3, width, label=labels[2], color=colors[2])

        plt.title(f"{title}\n", fontsize=15, fontweight="bold", pad=12)
        plt.text(0, 1.02, unit_label, transform=plt.gca().transAxes, ha='left', fontsize=10, color='gray')
        plt.xticks(x, years)
        plt.legend(loc='upper right')
        plt.gca().yaxis.set_major_formatter(ticker.FuncFormatter(lambda x, p: format(int(x), ',')))
        plt.grid(axis="y", linestyle="--", alpha=0.3)
        plt.tight_layout()
        
        save_path = os.path.join(target_dir, f"{file_name}_{suffix}.png")
        plt.savefig(save_path, dpi=150)
        plt.close()
        return save_path
    except Exception as e:
        print(f"  [Graph Error] {title}: {e}")
        return None

# =========================================================
# 4. Word 보고서 작성 (산업분야 추가)
# =========================================================
def save_as_word_report(data, file_name, target_dir):
    if not data: return None
    if not os.path.exists(target_dir): os.makedirs(target_dir)
    report_dir = target_dir
    doc = Document()
    
    header = data.get("Report_Header") or {}
    comp_name = header.get("Company_Name", file_name)
    ceo_name = header.get("CEO_Name", "확인 필요")
    rating = header.get("Investment_Rating", "-")
    # ✅ 산업분야 추출
    industry = header.get("Industry_Classification", "확인 필요")

    doc.add_paragraph("\n\n\n")
    t = doc.add_heading(f"{comp_name}\n투자 검토 보고서", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n" * 5)
    p = doc.add_paragraph(f"Date: 2025. XX. XX\nAnalyst: LUCEN Investment Intelligence")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_page_break()

    # [1] Executive Summary (산업분야 Row 추가)
    doc.add_heading("Executive Summary", level=1)
    
    # Row 3 -> 4로 변경
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    
    table.rows[0].cells[0].text = "기업명"
    table.rows[0].cells[1].text = str(comp_name)
    table.rows[1].cells[0].text = "대표자"
    table.rows[1].cells[1].text = str(ceo_name)
    
    # ✅ 산업분야 Row 추가
    table.rows[2].cells[0].text = "산업 분야"
    table.rows[2].cells[1].text = str(industry)
    
    table.rows[3].cells[0].text = "투자 등급"
    table.rows[3].cells[1].text = str(rating)
    
    apply_table_style(table)
    doc.add_paragraph("\n")
    doc.add_paragraph(data.get("Investment_Thesis_Summary") or "요약 내용 없음")

    # [2] 재무 현황
    doc.add_heading("1. 재무 현황 및 투자 유치 현황", level=1)
    fin = data.get("Financial_Status") or {}

    # 1-1. 재무상태표
    doc.add_heading("1-1. 재무상태표 (Balance Sheet)", level=2)
    bs = fin.get("Detailed_Balance_Sheet") or {}
    years = bs.get("Years") or []
    
    if years and len(years) > 0:
        row_keys = [
            ("유동자산", "Current_Assets"), ("비유동자산", "Non_Current_Assets"), ("자산총계", "Total_Assets"),
            ("유동부채", "Current_Liabilities"), ("비유동부채", "Non_Current_Liabilities"), ("부채총계", "Total_Liabilities"),
            ("자본금", "Capital_Stock"), ("이익잉여금외", "Retained_Earnings_Etc"), ("자본총계", "Total_Equity")
        ]
        t_bs = doc.add_table(rows=len(row_keys)+1, cols=len(years)+1)
        t_bs.style = "Table Grid"
        t_bs.rows[0].cells[0].text = "구분"
        for i, y in enumerate(years):
            t_bs.rows[0].cells[i+1].text = str(y)
        for r_idx, (label, key) in enumerate(row_keys):
            vals = bs.get(key) or []
            t_bs.rows[r_idx+1].cells[0].text = label
            for c_idx, val in enumerate(vals):
                if c_idx < len(years):
                    t_bs.rows[r_idx+1].cells[c_idx+1].text = str(val)
        apply_table_style(t_bs, header_color="FFFFCC")

        t_assets = bs.get("Total_Assets") or []
        t_liab = bs.get("Total_Liabilities") or []
        t_equity = bs.get("Total_Equity") or []
        if t_assets:
            doc.add_paragraph("\n■ 재무 상태 추이")
            bs_img = create_grouped_bar_chart_3(
                years, t_assets, t_liab, t_equity,
                ["자산총계", "부채총계", "자본총계"],
                [COLOR_BLUE, COLOR_YELLOW, COLOR_RED],
                "재무상태 변동 추이", file_name, "bs_chart", target_dir, 
                unit_label="(단위: 백만원)"
            )
            if bs_img: doc.add_picture(bs_img, width=Inches(6.0))
    else:
        doc.add_paragraph("재무상태표 데이터가 확인되지 않았습니다.")

    # 1-2. 손익계산서
    doc.add_heading("1-2. 수익성 분석 (Income Statement)", level=2)
    doc.add_paragraph(fin.get("Key_Financial_Commentary") or "")
    
    is_data = fin.get("Income_Statement_Summary") or {}
    is_years = is_data.get("Years") or []
    rev = is_data.get("Total_Revenue") or []
    op_prof = is_data.get("Operating_Profit") or []
    net_prof = is_data.get("Net_Profit") or []

    if is_years and len(is_years) > 0:
        doc.add_paragraph("\n■ 수익성 추이")
        is_img = create_grouped_bar_chart_3(
            is_years, rev, op_prof, net_prof,
            ["매출액", "영업이익", "당기순이익"],
            [COLOR_BLUE, COLOR_YELLOW, COLOR_RED],
            "수익성 변동 추이", file_name, "is_chart", target_dir,
            unit_label="(단위: 백만원)"
        )
        if is_img: doc.add_picture(is_img, width=Inches(6.0))
    else:
        doc.add_paragraph("손익계산서 데이터가 확인되지 않았습니다.")

    # 1-3. 투자 유치
    doc.add_heading("1-3. 투자 유치 현황", level=2)
    inv_hist = fin.get("Investment_History") or []
    if inv_hist and len(inv_hist) > 0:
        t_inv = doc.add_table(rows=1, cols=4)
        t_inv.style = "Table Grid"
        t_inv.rows[0].cells[0].text = "Date"
        t_inv.rows[0].cells[1].text = "Round"
        t_inv.rows[0].cells[2].text = "Amount"
        t_inv.rows[0].cells[3].text = "Investor"
        apply_table_style(t_inv)
        for item in inv_hist:
            row = t_inv.add_row().cells
            row[0].text = str(item.get("Date") or "")
            row[1].text = str(item.get("Round") or "")
            row[2].text = str(item.get("Amount") or "")
            row[3].text = str(item.get("Investor") or "")
    else:
        doc.add_paragraph("투자 이력이 확인되지 않았습니다.")
    
    # 1-4. 미래 수익 구조
    doc.add_heading("1-4. 미래 수익 구조", level=2)
    # [핵심 수정] fin_status 변수가 정의되지 않았을 경우를 대비해 여기서 다시 정의합니다.
    # 함수 인자로 받은 전체 데이터 변수명이 'data' 인지 'json_data' 인지 확인하세요.
    # 아래 코드는 Traceback에 나온 변수명인 'data'를 기준으로 작성되었습니다.
    # 만약 함수 인자가 json_data라면 data를 json_data로 변경해주세요.
    
    fin_status = data.get("Financial_Status", {}) 

    future_rev = fin_status.get("Future_Revenue_Structure", {})
    
    # 데이터 가져오기 (없을 경우 기본 문구 출력)
    bm_text = future_rev.get("Business_Model")
    if not bm_text:
        bm_text = "분석된 비즈니스 모델 내용이 없습니다."

    cash_cow_text = future_rev.get("Future_Cash_Cow")
    if not cash_cow_text:
        cash_cow_text = "분석된 Cash Cow 내용이 없습니다."
    
    # 문서에 쓰기
    p_bm = doc.add_paragraph()
    p_bm.add_run("■ 비즈니스 모델(BM) 및 수익 구조\n").bold = True
    p_bm.add_run(bm_text)
    
    doc.add_paragraph("") # 공백 추가
    
    p_cow = doc.add_paragraph()
    p_cow.add_run("■ 향후 Cash Cow 및 이익 기여도\n").bold = True
    p_cow.add_run(cash_cow_text)

    doc.add_page_break()

    # [3] 성장 가능성
    doc.add_heading("2. 성장 가능성", level=1)
    growth = data.get("Growth_Potential") or {}

    doc.add_heading("2-1. 시장 동향 및 레퍼런스", level=2)
    trends = growth.get("Target_Market_Trends") or []
    if trends:
        for t_item in trends:
            p = doc.add_paragraph()
            p.add_run(f"[{t_item.get('Type','Info')}] ").bold = True
            p.add_run(f"{t_item.get('Content','')}\n- Source: {t_item.get('Source','')}")
    else:
        doc.add_paragraph("관련 시장 동향 자료가 부족합니다.")

    # -----------------------------------------------------
    # [UPDATE] 2. 시장성 및 성장 잠재력
    # -----------------------------------------------------
    doc.add_heading("2. 시장성 및 성장 잠재력", level=1)
    
    # [FIX] 변수명을 'json_data' -> 'data'로 수정 (함수 인자와 일치)
    gp = data.get("Growth_Potential", {})

    # 2-1. 타겟 시장 분석 (New)
    doc.add_heading("2-1. 타겟 시장 분석 (Target Market Analysis)", level=2)
    market_analysis = gp.get("Target_Market_Analysis", {})
    
    # 3단 구성 출력 (Target Area / Characteristics / Positioning)
    if market_analysis:
        p_ma = doc.add_paragraph()
        p_ma.add_run("■ 타겟 영역 정의: ").bold = True
        p_ma.add_run(f"{market_analysis.get('Target_Area', '-')}\n")
        
        p_ma.add_run("■ 시장 특성: ").bold = True
        p_ma.add_run(f"{market_analysis.get('Market_Characteristics', '-')}\n")
        
        p_ma.add_run("■ 경쟁 포지셔닝: ").bold = True
        p_ma.add_run(f"{market_analysis.get('Competitive_Positioning', '-')}")
    else:
        doc.add_paragraph("타겟 시장 분석 데이터가 없습니다.")

    doc.add_paragraph("") # Spacer

    # 2-2. 시장 동향 및 레퍼런스 (기존 항목)
    doc.add_heading("2-2. 시장 동향 및 레퍼런스 (Market Trends)", level=2)
    trends = gp.get("Target_Market_Trends", [])
    if trends:
        for t in trends:
            # Type과 Source가 있는 경우와 없는 경우를 깔끔하게 처리
            source_text = f" (출처: {t.get('Source')})" if t.get('Source') else ""
            type_text = f"[{t.get('Type')}] " if t.get('Type') else ""
            doc.add_paragraph(f"• {type_text}{t.get('Content', '')}{source_text}", style='List Bullet')
    else:
        doc.add_paragraph("관련 시장 동향 데이터가 확인되지 않았습니다.")

    doc.add_paragraph("") # Spacer

    # 2-3. L/O(기술이전) 및 Exit 전략 (New)
    doc.add_heading("2-3. L/O(기술이전) 및 Exit 전략", level=2)
    lo_strat = gp.get("LO_Exit_Strategy", {})
    
    # (1) 이미 검증된 시그널
    doc.add_paragraph("■ 이미 검증된 시그널 (레퍼런스)", style='List Paragraph')
    signals = lo_strat.get("Verified_Signals", [])
    if signals:
        for sig in signals:
            doc.add_paragraph(f"- {sig}", style='List Bullet 2')
    else:
        doc.add_paragraph("- 확인된 시그널 없음", style='List Bullet 2')

    # (2) 예상 L/O 시나리오 (Table)
    doc.add_paragraph("\n■ 예상 L/O 시나리오", style='List Paragraph')
    scenarios = lo_strat.get("Expected_LO_Scenarios", [])
    if scenarios:
        # 표 생성 (Header + Data)
        lo_table = doc.add_table(rows=1, cols=3)
        lo_table.style = 'Table Grid'
        
        # Header 설정
        hdr = lo_table.rows[0].cells
        hdr[0].text = "구분"
        hdr[1].text = "가능성"
        hdr[2].text = "코멘트"
        
        # Data 입력
        for sc in scenarios:
            row_cells = lo_table.add_row().cells
            row_cells[0].text = sc.get("Category", "-")
            row_cells[1].text = sc.get("Probability", "-")
            row_cells[2].text = sc.get("Comment", "-")
    else:
        doc.add_paragraph("제시된 시나리오가 없습니다.")

    # (3) 적정 가치 범위
    doc.add_paragraph("")
    p_val = doc.add_paragraph()
    p_val.add_run("■ 적정 가치 범위 (보수적 판단): ").bold = True
    p_val.add_run(lo_strat.get("Valuation_Range", "산정 불가"))

    doc.add_paragraph("") # Spacer

    # -----------------------------------------------------
    # 2-4. 주요 성장 지표 (Growth Indicators) - 그래프 복원
    # -----------------------------------------------------
    doc.add_heading("2-4. 주요 성장 지표 (Growth Indicators)", level=2)
    
    # gp는 함수 상단에서 정의된 data.get("Growth_Potential", {}) 입니다.
    stats = gp.get("Export_and_Contract_Stats", {})

    # 1. 수출 실적 그래프
    if stats.get("Export_Graph_Data"):
        doc.add_paragraph("■ 수출 실적")
        # report_dir는 함수 인자로 받은 저장 경로입니다.
        img = create_basic_bar_chart(
            stats["Export_Graph_Data"], 
            "수출 실적 추이", 
            COLOR_RED, 
            file_name, 
            "export", 
            report_dir
        )
        if img: 
            doc.add_picture(img, width=Inches(5.5))
            doc.add_paragraph("") # 간격

    # 2. 계약 건수 그래프
    if stats.get("Contract_Count_Graph_Data"):
        doc.add_paragraph("■ 계약 건수")
        img = create_basic_bar_chart(
            stats["Contract_Count_Graph_Data"], 
            "계약 체결 추이", 
            COLOR_YELLOW, 
            file_name, 
            "contract", 
            report_dir
        )
        if img: 
            doc.add_picture(img, width=Inches(5.5))
            doc.add_paragraph("") # 간격
            
    # 3. 매출 성장 그래프
    if stats.get("Sales_Graph_Data"):
        doc.add_paragraph("■ 매출 규모(성장성 관점)")
        img = create_basic_bar_chart(
            stats["Sales_Graph_Data"], 
            "매출 성장 추이", 
            COLOR_BLUE, 
            file_name, 
            "sales_growth", 
            report_dir
        )
        if img: 
            doc.add_picture(img, width=Inches(5.5))
            doc.add_paragraph("") # 간격

    # 데이터가 아예 없는 경우 멘트 처리
    if not (stats.get("Export_Graph_Data") or stats.get("Contract_Count_Graph_Data") or stats.get("Sales_Graph_Data")):
        doc.add_paragraph("시각화할 주요 성장 지표 데이터(그래프)가 부족합니다.")

    doc.add_page_break()

    # -----------------------------------------------------
    # [UPDATE] 3. 기술 경쟁력 및 파이프라인
    # -----------------------------------------------------
    doc.add_heading("3. 기술 경쟁력 및 파이프라인", level=1)
    
    # JSON Schema 변경에 맞춰 데이터 키 변경
    tech_data = data.get("Technology_and_Pipeline", {})

    # 3-1. Market Pain Points
    doc.add_heading("3-1. Market Pain Points", level=2)
    pain_points = tech_data.get("Market_Pain_Points", [])
    
    if pain_points:
        for pp in pain_points:
            doc.add_paragraph(f"• {pp}", style='List Bullet')
    else:
        doc.add_paragraph("분석된 시장 문제점(Pain Points) 데이터가 없습니다.")

    doc.add_paragraph("") # Spacer

    # 3-2. Solution & Core Tech
    doc.add_heading("3-2. Solution & Core Tech", level=2)
    sol_data = tech_data.get("Solution_and_Core_Tech", {})
    
    tech_name = sol_data.get("Technology_Name", "핵심 기술명 미확인")
    p_tech = doc.add_paragraph()
    p_tech.add_run(f"■ 핵심 솔루션: {tech_name}\n").bold = True
    
    features = sol_data.get("Key_Features", [])
    if features:
        for ft in features:
            doc.add_paragraph(f"- {ft}", style='List Bullet 2')
    else:
        doc.add_paragraph("기술 상세 특징 데이터가 없습니다.", style='List Bullet 2')

    doc.add_paragraph("") # Spacer

    # 3-3. 주요 파이프라인 개발 현황 (New)
    doc.add_heading("3-3. 주요 파이프라인 개발 현황", level=2)
    pipeline_status = tech_data.get("Pipeline_Development_Status", {})

    # (1) 핵심 플랫폼 기술 상세
    doc.add_paragraph("■ 핵심 플랫폼 기술 상세", style='List Paragraph')
    platform_details = pipeline_status.get("Core_Platform_Details", "내용 없음")
    doc.add_paragraph(platform_details)
    doc.add_paragraph("")

    # (2) 기술적 위험도(Risk) 분석
    doc.add_paragraph("■ 기술적 위험도(Risk) 분석", style='List Paragraph')
    risk_analysis = pipeline_status.get("Technical_Risk_Analysis", "내용 없음")
    doc.add_paragraph(risk_analysis)
    doc.add_paragraph("")

    # (3) 기술성 결론
    doc.add_paragraph("■ 기술성 결론", style='List Paragraph')
    tech_conclusion = pipeline_status.get("Technical_Conclusion", "내용 없음")
    
    # 결론 부분을 강조하기 위해 배경색이나 테두리를 주면 좋겠지만, 
    # python-docx 기본 기능으로는 굵게 처리하거나 인용구 스타일을 쓰는 것이 안전합니다.
    p_conc = doc.add_paragraph()
    runner = p_conc.add_run(tech_conclusion)
    runner.bold = False 
    
    doc.add_page_break()

    # [Section 4] 주요 인력 및 조직 (✅ 대폭 수정됨)
    doc.add_heading("4. 주요 인력 및 조직", level=1)
    # [수정] Processor의 새로운 키 Key_Personnel에 맞춤
    kp = data.get("Key_Personnel") or {}
    
    doc.add_heading("4-1. 대표이사 레퍼런스", level=2)
    ceo = kp.get("CEO_Reference") or {}
    doc.add_paragraph(f"■ 성명: {ceo.get('Name', '')}")
    doc.add_paragraph(f"■ 학력 및 경력:\n{ceo.get('Background_and_Education', '')}")
    doc.add_paragraph(f"■ 핵심 역량:\n{ceo.get('Core_Competency', '')}")
    doc.add_paragraph(f"■ 경영 철학:\n{ceo.get('Management_Philosophy', '')}")
    doc.add_paragraph(f"■ VC 관점 평가:\n{ceo.get('VC_Perspective_Evaluation', '')}")

    doc.add_heading("4-2. 조직 역량", level=2)
    team = kp.get("Team_Capability") or {}
    doc.add_paragraph("■ 핵심 임원진:")
    for ex in team.get("Key_Executives") or []: doc.add_paragraph(f"- {ex}")
    doc.add_paragraph(f"■ 조직 강점:\n{team.get('Organization_Strengths', '')}")
    doc.add_paragraph(f"■ 자문단:\n{team.get('Advisory_Board', '')}")
    
    doc.add_page_break()

    # -----------------------------------------------------
    # [UPDATE] 5. 리스크 및 종합 투자 판단
    # -----------------------------------------------------
    doc.add_heading("5. 리스크 및 종합 투자 판단", level=1)
    
    # [변경됨] 1. 리스크 데이터는 이제 최상위 키 "Key_Risks_and_Mitigation"에 있습니다.
    risks = data.get("Key_Risks_and_Mitigation") or []

    # 5-1. 주요 리스크 및 대응 (Key Risks)
    doc.add_heading("5-1. 주요 리스크 및 대응 (Key Risks)", level=2)
    
    if risks:
        rt = doc.add_table(rows=1, cols=2)
        rt.style = "Table Grid"
        rt.rows[0].cells[0].text = "Risk Factor"
        rt.rows[0].cells[1].text = "Mitigation Strategy"
        
        try:
            apply_table_style(rt)
        except NameError:
            pass 
            
        for r in risks:
            row = rt.add_row().cells
            row[0].text = str(r.get("Risk_Factor") or "")
            row[1].text = str(r.get("Mitigation_Strategy") or "")
    else:
        doc.add_paragraph("식별된 주요 리스크 데이터가 없습니다.")

    # [변경됨] 2. 밸류에이션과 판단은 새로운 키 "Valuation_and_Judgment"에 묶여 있습니다.
    val_judge_data = data.get("Valuation_and_Judgment") or {}

    # 5-2. 밸류에이션 추정
    doc.add_heading("5-2. 밸류에이션 추정 (Valuation Estimation)", level=2)
    
    # 1. 밸류에이션 테이블 (기존)
    val_table = val_judge_data.get("Valuation_Table") or []
    if val_table:
        vt = doc.add_table(rows=1, cols=4)
        vt.style = "Table Grid"
        vt.rows[0].cells[0].text = "Round"
        vt.rows[0].cells[1].text = "Pre-Money"
        vt.rows[0].cells[2].text = "Post-Money"
        vt.rows[0].cells[3].text = "Comment"
        apply_table_style(vt)
        for v in val_table:
            r = vt.add_row().cells
            r[0].text = str(v.get("Round") or "-")
            r[1].text = str(v.get("Pre_Money") or "-")
            r[2].text = str(v.get("Post_Money") or "-")
            r[3].text = str(v.get("Comment") or "-")
    
    # 2. 밸류에이션 산정 로직 (NEW - 상세 출력)
    logic = val_judge_data.get("Valuation_Logic_Detail") or {}
    if logic:
        doc.add_paragraph("") # Spacer
        p_logic = doc.add_paragraph()
        p_logic.add_run("■ 밸류에이션 산정 로직 (Analyst Opinion)\n").bold = True
        
        peers = ", ".join(logic.get("Peer_Group") or [])
        p_logic.add_run(f"• 비교 기업(Peer Group): {peers}\n")
        p_logic.add_run(f"• 적용 지표: {logic.get('Applied_Multiple', '-')}\n")
        p_logic.add_run(f"• 적용 이익: {logic.get('Target_Net_Income', '-')}\n")
        p_logic.add_run(f"• 산출 근거:\n{logic.get('Calculation_Rationale', '-')}")
    else:
        doc.add_paragraph("상세 밸류에이션 로직 데이터가 없습니다.")

    # 5-3. 종합 투자 판단 (VC/AC 시각)
    doc.add_heading("5-3. 종합 투자 판단 (VC/AC 시각)", level=2)
    
    # (1) 3대 평가 축
    doc.add_paragraph("■ 3대 평가 축 분석")
    axes = val_judge_data.get("Three_Axis_Assessment") or {}
    
    if axes:
        doc.add_paragraph(f"- 기술성: {axes.get('Technology_Rating') or '-'}", style='List Bullet')
        doc.add_paragraph(f"- 성장성: {axes.get('Growth_Rating') or '-'}", style='List Bullet')
        doc.add_paragraph(f"- 회수 가능성: {axes.get('Exit_Rating') or '-'}", style='List Bullet')
    else:
        doc.add_paragraph("평가 축 분석 데이터가 없습니다.")

    # (2) 적합한 투자자 유형
    doc.add_paragraph("")
    doc.add_paragraph("■ 적합한 투자자 유형")
    doc.add_paragraph(str(val_judge_data.get("Suitable_Investor_Type") or "분석 불가"))

    # [변경됨] 3. 최종 결론은 다시 최상위 키 "Final_Conclusion"에 있습니다.
    doc.add_heading("5-4. 종합 결론 (Final Conclusion)", level=2)
    doc.add_paragraph(str(data.get("Final_Conclusion") or "종합 결론 내용이 없습니다."))

    # [기존 코드 유지]
    out_path = os.path.join(target_dir, f"{file_name}_검토보고서.docx")
    doc.save(out_path)
    return out_path

# =========================================================
# 5. Main (Skip 로직 추가)
# =========================================================
def main():
    input_dir = "data"
    output_dir = "output"
    report_dir = "output_report"
    
    for d in [output_dir, report_dir]:
        os.makedirs(d, exist_ok=True)

    pdf_files = glob.glob(os.path.join(input_dir, "*.pdf"))
    if not pdf_files:
        print("PDF 파일이 없습니다.")
        return

    print(f"총 {len(pdf_files)}개 파일 처리 시작...")

    for i, file_path in enumerate(pdf_files):
        file_name = os.path.splitext(os.path.basename(file_path))[0]
        json_path = os.path.join(output_dir, f"{file_name}_refined.json")
        docx_path = os.path.join(report_dir, f"{file_name}_검토보고서.docx")
        
        print(f"\n>>> [{i+1}/{len(pdf_files)}] {file_name}")

        # ✅ Skip Logic: JSON과 DOCX가 모두 존재하면 건너뜀
        if os.path.exists(json_path) and os.path.exists(docx_path):
            print(f"   [Skip] 이미 완료된 파일입니다. ({file_name})")
            continue

        try:
            data = refine_pdf_to_json_onecall(file_path)
            
            if data is None or (isinstance(data, dict) and data.get("error")):
                err = data.get("error") if isinstance(data, dict) else "Data is None"
                print(f"   [Error] JSON 생성 실패: {err}")
                continue

            with open(json_path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)

            doc_path = save_as_word_report(data, file_name, report_dir)
            if doc_path:
                print(f"   [완료] {doc_path}")
            else:
                print("   [Skip] 보고서 생성 실패")

        except Exception as e:
            import traceback
            print(f"   [Fail] {e}")
            traceback.print_exc()

if __name__ == "__main__":
    main()