import os
import re
import io
import datetime
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
import numpy as np
from collections import Counter
from typing import Optional, Tuple

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# =========================================================
# ✅ [설정] Matplotlib 한글 폰트 등
# =========================================================
plt.rcParams["font.family"] = "Malgun Gothic"
plt.rcParams["axes.unicode_minus"] = False

COLOR_RED = "indianred"
COLOR_YELLOW = "khaki"
COLOR_BLUE = "cornflowerblue"

# =========================================================
# 1. 공통 유틸리티 & 문서 스타일링 헬퍼
# =========================================================
def set_cell_background(cell, fill_color):
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def apply_table_style(table, header_color="E7E6E6", center_all=True):
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if center_all:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i == 0: 
                set_cell_background(cell, header_color)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs: run.bold = True

def apply_table_colors(table, header_color="D9D9D9", first_col_color="F2F2F2"):
    if len(table.rows) == 0: return
    for cell in table.rows[0].cells:
        set_cell_background(cell, header_color)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs: run.bold = True
    for row in table.rows[1:]:
        cell = row.cells[0]
        set_cell_background(cell, first_col_color)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs: run.bold = True
            
def add_smart_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.keep_with_next = True
    return h

def add_formatted_text(paragraph, text):
    if not isinstance(text, str):
        paragraph.add_run(str(text))
        return
    parts = re.split(r'(<sup>|</sup>|<sub>|</sub>)', text)
    is_sup, is_sub = False, False
    for part in parts:
        if part.lower() == '<sup>': is_sup = True
        elif part.lower() == '</sup>': is_sup = False
        elif part.lower() == '<sub>': is_sub = True
        elif part.lower() == '</sub>': is_sub = False
        elif part: 
            run = paragraph.add_run(part)
            if is_sup: run.font.superscript = True
            if is_sub: run.font.subscript = True

# 🚨 내장되어 있던 함수들을 전역으로 분리
def format_financial_number(val):
    s = str(val).strip()
    if not s or s in ['-', 'N/A']: return s
    if "," in s: return s 
    is_bracket = False
    clean_s = s.replace(" ", "")
    if clean_s.startswith("(") and clean_s.endswith(")"):
        is_bracket = True
        clean_s = clean_s[1:-1]
    elif clean_s.startswith("△") or clean_s.startswith("▽"):
        clean_s = "-" + clean_s[1:]
    try:
        f = float(clean_s)
        formatted = f"{int(f):,}" if f.is_integer() else f"{f:,}"
        if is_bracket: return f"({formatted})"
        elif clean_s.startswith("-") and not formatted.startswith("-"): return f"-{formatted}"
        return formatted
    except ValueError:
        return s

def extract_date_from_filename(fname):
    match = re.search(r'(20[1-3]\d)[-._]?([0-1]\d)', fname)
    if match:
        year, month = match.group(1), match.group(2)
        if 1 <= int(month) <= 12: return f"{year}.{month}"
    return None

def detect_dynamic_unit(data_list, is_count=False, fallback_unit="백만원"):
    if is_count: return "건"
    units = []
    for val in data_list:
        val_str = str(val)
        if "백만" in val_str: units.append("백만원")
        elif "억" in val_str: units.append("억원")
        elif "조" in val_str: units.append("조원")
        elif "천" in val_str: units.append("천원")
        elif "만" in val_str: units.append("만원")
        elif "%" in val_str: units.append("%")
        elif "$" in val_str or "달러" in val_str: units.append("달러")
    if units: return Counter(units).most_common(1)[0][0]
    return fallback_unit

# =========================================================
# 2. 인메모리(In-Memory) 그래프 생성 함수 (속도 개선)
# =========================================================
def create_basic_bar_chart_stream(data_list, title, color, fallback_unit="백만원"):
    """디스크 저장 없이 RAM에서 바로 이미지를 생성하여 docx에 전달하는 함수"""
    if not data_list or len(data_list) < 2: return None
    current_year = datetime.datetime.now().year
    is_count = "건" in title or "계약" in title
    raw_values = [row[1] for row in data_list[1:] if len(row) > 1]
    
    base_unit = detect_dynamic_unit(raw_values, is_count, fallback_unit)
    unit_label = f"(단위: {base_unit})"
    short_unit = "억" if base_unit == "억원" else ("건" if is_count else "")
    
    try:
        years, values = [], []
        for row in data_list[1:]:
            if isinstance(row, list) and len(row) >= 2:
                y = str(row[0]).replace("년","").strip()
                year_match = re.search(r'\d{4}', y)
                if year_match:
                    yr = int(year_match.group())
                    y = f"{yr}(E)" if "E" in y.upper() or yr >= current_year else str(yr)
                
                raw_val = str(row[1]).strip()
                num_str = re.sub(r'[^\d\.\-]', '', raw_val)
                v = 0.0
                if num_str:
                    v = float(num_str)
                    if base_unit == "억원" and "백만" in raw_val: v /= 100
                    elif base_unit == "백만원" and "억" in raw_val: v *= 100
                years.append(y)
                values.append(v)
                
        if not years or sum(values) == 0: return None

        plt.figure(figsize=(7, 4.5))
        bars = plt.bar(years, values, color=color, width=0.5, edgecolor='dimgray', alpha=0.85)
        plt.title(f"{title}\n", fontsize=14, fontweight="bold", pad=10)
        plt.text(0, 1.02, unit_label, transform=plt.gca().transAxes, ha='left', fontsize=10, color='gray')
        plt.gca().yaxis.set_major_formatter(ticker.FuncFormatter(lambda x, pos: format(int(x), ',')))
        plt.grid(axis="y", linestyle="--", alpha=0.3)
        max_val = max(values) if values else 1
        plt.ylim(bottom=0, top=max_val * 1.25) 
        
        for bar in bars:
            height = bar.get_height()
            if height > 0:
                fmt_val = f"{int(height):,}" if height.is_integer() else f"{height:,.1f}"
                plt.text(bar.get_x() + bar.get_width()/2.0, height + (max_val*0.02),
                         f"{fmt_val}{short_unit}", ha='center', va='bottom', fontsize=9, fontweight='bold', color='#1F497D')

        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150)
        plt.close()
        buf.seek(0)
        return buf # 파일 경로 대신 버퍼 반환
    except Exception as e:
        print(f"  [Graph Error] {title}: {e}")
        return None

def create_investment_bubble_chart_stream(sorted_inv):
    """투자 이력 버블 차트 인메모리 생성 모듈화 (범위 데이터 예외처리 완벽 적용)"""
    try:
        x_positions, x_labels, y_amounts, rounds, sizes, bubble_colors, amount_labels = [], [], [], [], [], [], []
        def get_pastel_color(r_name):
            r = str(r_name).lower().replace(" ", "").replace("-", "")
            if "seed" in r or "시드" in r: return "#FDFD96"
            elif "pre" in r or "프리" in r: return "#77DD77"
            elif "seriesa" in r or "시리즈a" in r: return "#FF6961"
            elif "seriesb" in r or "시리즈b" in r: return "#AEC6CF"
            elif "seriesc" in r or "시리즈c" in r: return "#C3B1E1"
            elif "m&a" in r or "ipo" in r or "상장" in r or "인수" in r: return "#77DD77"
            return "#E5E4E2"

        for i, item in enumerate(sorted_inv):
            round_txt = str(item.get("Round", "")).strip()
            if not round_txt or round_txt.lower() in ["none", "null", "비공개", "-"]: round_txt = "N/A"

            raw_date = str(item.get("Date", "")).strip()
            year_match = re.search(r'\d{4}', raw_date)
            display_date = year_match.group(0) if year_match else round_txt 
            
            raw_amt = str(item.get("Amount", "")).strip()
            
            # 🚨 [핵심 수정] 금액 문자열에서 숫자들을 모두 찾아 배열로 추출 ("60~100" -> ['60', '100'])
            nums = re.findall(r'\d+\.?\d*', raw_amt)
            
            if not nums:
                amt_val, amt_str = 5.0, "비공개" 
            else:
                # 1. 단위 파악 (배수 결정)
                multiplier = 1.0
                if "백만" in raw_amt: multiplier = 1/100.0
                elif "조" in raw_amt: multiplier = 10000.0
                elif "Mil" in raw_amt or "M" in raw_amt or "달러" in raw_amt: multiplier = 13.0
                
                # 2. 추출된 숫자들에 각각 배수 적용
                scaled_nums = [float(n) * multiplier for n in nums]
                
                # 3. Y축 좌표 및 라벨 동적 생성
                if len(scaled_nums) >= 2:
                    # 범위인 경우 ("60~100억") -> Y축은 평균값(80)으로, 라벨은 "60~100억" 유지
                    amt_val = (scaled_nums[0] + scaled_nums[1]) / 2.0
                    amt_str = f"{scaled_nums[0]:g}~{scaled_nums[1]:g}억"
                else:
                    # 단일 값인 경우 ("50억")
                    amt_val = scaled_nums[0]
                    amt_str = f"{amt_val:g}억"
                    
            x_positions.append(i) 
            x_labels.append(display_date)
            y_amounts.append(amt_val)
            rounds.append(round_txt)
            amount_labels.append(amt_str)
            sizes.append(max(1000, amt_val * 25))
            bubble_colors.append(get_pastel_color(round_txt))
            
        if not x_positions: return None
        
        plt.figure(figsize=(7, 3.5))
        plt.scatter(x_positions, y_amounts, s=sizes, c=bubble_colors, alpha=0.9, edgecolors='dimgray', linewidth=1.5)
        for idx in range(len(x_positions)):
            plt.annotate(rounds[idx], (x_positions[idx], y_amounts[idx]), ha='center', va='center', fontsize=9, fontweight='bold', color='black')
            radius_pt = (sizes[idx] ** 0.5) / 2
            plt.annotate(amount_labels[idx], (x_positions[idx], y_amounts[idx]), xytext=(0, radius_pt + 3), textcoords="offset points", ha='center', va='bottom', fontsize=9, fontweight='bold', color='#1F497D')
        
        plt.xticks(x_positions, x_labels)
        plt.xlim(left=-0.5, right=len(x_positions) - 0.5)
        max_amt = max(y_amounts) if y_amounts else 10
        plt.ylim(bottom=0, top=max_amt * 1.5) 
        plt.ylabel("투자 금액 (억원)")
        plt.title("Investment Timeline")
        plt.grid(axis='y', linestyle='--', alpha=0.3)
        plt.tight_layout()
        
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=200, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf
    except Exception as e:
        print(f"  [Bubble Chart Error]: {e}")
        return None

# =========================================================
# 3. 워드 섹션별 렌더링 로직 (모듈화)
# =========================================================
def _render_executive_summary(doc, data, header, fin_data):
    doc.add_heading("Executive Summary", level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "기업명"
    table.rows[0].cells[1].text = str(header.get("Company_Name", ""))
    table.rows[1].cells[0].text = "대표자"
    table.rows[1].cells[1].text = str(header.get("CEO_Name", ""))
    table.rows[2].cells[0].text = "산업 분야"
    table.rows[2].cells[1].text = str(header.get("Industry_Classification", ""))
    table.rows[3].cells[0].text = "투자 등급"
    table.rows[3].cells[1].text = str(header.get("Investment_Rating", "평가 대기"))
    
    try: apply_table_colors(table, header_color="D9D9D9", first_col_color="F2F2F2")
    except NameError: apply_table_style(table)
        
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")

    highlights = data.get("Investment_Highlights") or fin_data.get("Investment_Highlights") or []
    if highlights and isinstance(highlights, list):
        for item in highlights:
            title, logic = str(item.get("Highlight_Title", "")).strip(), str(item.get("Highlight_Logic", "")).strip()
            if title and not title.startswith("["): title = f"[{title}]"
            p_summary = doc.add_paragraph(style='List Bullet')
            p_summary.add_run(f"{title} ").bold = True
            p_summary.add_run(logic)
    else:
        legacy = data.get("Investment_Thesis_Summary") or fin_data.get("Investment_Thesis_Summary")
        doc.add_paragraph(legacy if legacy else "요약 정보가 없습니다.")
    doc.add_paragraph("\n")

def _render_financial_and_investment(doc, data, file_name, original_file_name, fin_data):
    doc.add_heading("1. 재무 현황 및 투자 유치 현황", level=1)
    
    # 1-1. 재무상태표
    doc.add_heading("1-1. 재무상태표 (Balance Sheet)", level=2)
    bs_data = fin_data.get("Balance_Sheet", {})
    if isinstance(bs_data, list): bs_data = {}
    bs_cols, bs_rows = bs_data.get("Columns", []), bs_data.get("Rows", [])
    
    if bs_cols and bs_rows:
        doc.add_paragraph(bs_data.get("Unit", "단위 : 백만원")).alignment = WD_ALIGN_PARAGRAPH.RIGHT
        t_bs = doc.add_table(rows=1, cols=len(bs_cols))
        t_bs.style = "Table Grid"
        for i, col_name in enumerate(bs_cols): t_bs.rows[0].cells[i].text = str(col_name)
        for row_data in bs_rows:
            row_cells = t_bs.add_row().cells
            for i, val in enumerate(row_data):
                if i < len(row_cells): row_cells[i].text = format_financial_number(val)
        apply_table_colors(t_bs, header_color="D9D9D9", first_col_color="F2F2F2")
        for row in t_bs.rows:
            for cell in row.cells:
                for p in cell.paragraphs: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("IR 자료 내에 과거 재무상태표(자산/자본/부채) 데이터가 명시되어 있지 않습니다.")
    doc.add_paragraph("\n")

    # 1-2. 추정손익계산서
    is_data = fin_data.get("Income_Statement", {})
    if isinstance(is_data, list): is_data = {}
    is_cols, is_rows = is_data.get("Columns", []), is_data.get("Rows", [])
    current_year = datetime.datetime.now().year
    
    if is_cols:
        for i in range(1, len(is_cols)):
            match = re.search(r'\d{4}', str(is_cols[i]))
            if match:
                yr = int(match.group())
                is_cols[i] = f"{yr}(E)" if "E" in str(is_cols[i]).upper() or yr >= current_year else str(yr)

    year_count = len(is_cols) - 1 if len(is_cols) > 1 else 0
    doc.add_heading(f"1-2. {year_count if year_count > 0 else 5}개년 추정손익 (Income Statement)", level=2)
    
    if is_cols and is_rows:
        target_name = original_file_name if original_file_name else file_name
        ir_date_str = extract_date_from_filename(target_name)
        current_date_str = datetime.datetime.now().strftime("%Y-%m-%d")
        warn_txt = f"※ 본 정보는 {current_date_str} 분석 시점 기준이며, 원본 IR 자료의 작성 시기{f'({ir_date_str})' if ir_date_str else ''}에 따라 이미 경과된 연도임에도 추정치(E)로 표기되어 있을 수 있습니다."
        
        run_warn = doc.add_paragraph().add_run(warn_txt)
        run_warn.font.size, run_warn.font.color.rgb = Pt(8.5), RGBColor(128, 128, 128)
        doc.add_paragraph(is_data.get("Unit", "단위 : 백만원")).alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
        t_is = doc.add_table(rows=1, cols=len(is_cols))
        t_is.style = "Table Grid"
        for i, col_name in enumerate(is_cols): t_is.rows[0].cells[i].text = str(col_name)
        for row_data in is_rows:
            row_cells = t_is.add_row().cells
            for i, val in enumerate(row_data):
                if i < len(row_cells): row_cells[i].text = format_financial_number(val)
        apply_table_colors(t_is, header_color="D9D9D9", first_col_color="F2F2F2")
        for row in t_is.rows:
            for cell in row.cells:
                for p in cell.paragraphs: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("추정손익 데이터가 제공되지 않았습니다.")
    doc.add_paragraph("\n")

    # 1-3. 투자 유치 현황
    doc.add_heading("1-3. 투자 유치 현황", level=2)
    inv = fin_data.get("Investment_History") or []
    
    if inv:
        for item in inv:
            raw_date, round_txt = str(item.get("Date", "")).strip(), str(item.get("Round", "")).strip()
            if not re.search(r'\d{4}', raw_date) or "2000" in raw_date or "0000" in raw_date or raw_date.lower() in ["none", "null", "-", "n/a", ""]:
                item["Date"] = round_txt if round_txt else "-"

        ROUND_ORDER = {"angel":1,"seed":2,"prea":3,"seriesa":4,"preb":5,"seriesb":6,"prec":7,"seriesc":8,"preipo":9,"ipo":10, "엔젤":1,"시드":2,"프리a":3,"시리즈a":4,"프리b":5,"시리즈b":6,"프리c":7,"시리즈c":8,"프리ipo":9,"상장":10}
        def get_sort_key(item):
            raw_date = str(item.get("Date", ""))
            year_match = re.search(r'\d{4}', raw_date)
            year = int(year_match.group()) if year_match else 9999
            r_name = str(item.get("Round", "")).lower().replace(" ", "").replace("-", "")
            return (year, next((v for k, v in ROUND_ORDER.items() if k in r_name), 99))

        sorted_inv = sorted(inv, key=get_sort_key)

        t_inv = doc.add_table(rows=1, cols=4)
        t_inv.style = "Table Grid"
        for i, h in enumerate(["Date", "Round", "Amount", "Investor"]): t_inv.rows[0].cells[i].text = h
        for item in sorted_inv:
            row = t_inv.add_row().cells
            row[0].text, row[1].text, row[2].text, row[3].text = str(item.get("Date") or "-"), str(item.get("Round") or "-"), str(item.get("Amount") or "-"), str(item.get("Investor") or "-")
        apply_table_colors(t_inv, header_color="D9D9D9", first_col_color="F2F2F2")
        for row in t_inv.rows:
            for cell in row.cells:
                for p in cell.paragraphs: p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 🚨 메모리 버퍼 차트 삽입 (디스크 저장 제거)
        chart_buf = create_investment_bubble_chart_stream(sorted_inv)
        if chart_buf:
            doc.add_paragraph()
            doc.add_picture(chart_buf, width=Inches(6.0))
    else:
        doc.add_paragraph("투자 유치 이력 없음")

def _render_market_and_growth(doc, data):
    doc.add_heading("1-4. 미래 수익 구조", level=2)
    fr = data.get("Financial_Status", {}).get("Future_Revenue_Structure") or {}
    doc.add_paragraph(f"■ 비즈니스 모델: {fr.get('Business_Model', '내용 없음')}")
    doc.add_paragraph(f"■ 향후 Cash Cow: {fr.get('Future_Cash_Cow', '내용 없음')}")

    doc.add_heading("2. 시장성 및 성장 잠재력", level=1)
    mg = data.get("Growth_Potential") or data.get("Market_Analysis") or {}
    
    doc.add_heading("2-1. 타겟 시장 규모 (TAM-SAM-SOM)", level=2)
    p_market = doc.add_paragraph()
    target_analysis = mg.get("Target_Market_Analysis", {})
    add_formatted_text(p_market, f"• [TAM] Total Addressable Market: {target_analysis.get('TAM', '자료 부족')}\n")
    add_formatted_text(p_market, f"• [SAM] Serviceable Available Market: {target_analysis.get('SAM', '자료 부족')}\n")
    add_formatted_text(p_market, f"• [SOM] Serviceable Obtainable Market: {target_analysis.get('SOM', '자료 부족')}")
    
    doc.add_heading("2-2. 경쟁사 비교 분석", level=2)
    competitors = mg.get("Competitors_Comparison", [])
    if competitors and isinstance(competitors, list):
        t_comp = doc.add_table(rows=len(competitors)+1, cols=4)
        t_comp.style = "Table Grid"
        for i, h in enumerate(["기업명", "주요 제품/서비스", "타겟 시장", "핵심 기술 및 차별점"]): t_comp.rows[0].cells[i].text = h
        for i, comp in enumerate(competitors):
            r = t_comp.rows[i+1].cells
            r[0].text, r[1].text, r[2].text, r[3].text = str(comp.get("Company", "")), str(comp.get("Product", "")), str(comp.get("Target_Market", "")), str(comp.get("Core_Tech", ""))
        try: apply_table_colors(t_comp, header_color="D9D9D9", first_col_color="F2F2F2")
        except NameError: pass
        for i, row in enumerate(t_comp.rows):
            for j, cell in enumerate(row.cells):
                if i == 0 or j == 0:
                    for p in cell.paragraphs: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("\n")
    else:
        doc.add_paragraph("경쟁사 비교 데이터가 제공되지 않았습니다.\n")

    doc.add_heading("2-3. 주요 시장 트렌드", level=2)
    trends = mg.get("Target_Market_Trends", [])
    if trends:
        for t in trends: doc.add_paragraph(style='List Bullet').add_run(f"[{t.get('Type')}] {t.get('Content')} (Source: {t.get('Source')})")
    else: doc.add_paragraph("관련 시장 트렌드 정보가 없습니다.")
    doc.add_paragraph()

    # L/O 및 Exit
    doc.add_heading("2-4. L/O 및 Exit 전략", level=2)
    lo = mg.get("LO_Exit_Strategy") or {}
    add_formatted_text(doc.add_paragraph(), f"• 검증된 시그널: {', '.join(lo.get('Verified_Signals') or [])}")
    add_formatted_text(doc.add_paragraph(), f"• 적정 가치 범위: {lo.get('Valuation_Range', '확인 필요')}")
    
    raw_scenarios = lo.get("Expected_LO_Scenarios")
    if raw_scenarios:
        t_lo = doc.add_table(rows=1, cols=3)
        t_lo.style = "Table Grid"
        for i, h in enumerate(["구분", "가능성", "코멘트"]): t_lo.rows[0].cells[i].text = h
        try: apply_table_colors(t_lo, header_color="D9D9D9", first_col_color="F2F2F2")
        except NameError: apply_table_style(t_lo)
        
        normalized = []
        for s in raw_scenarios:
            if "M&A" in str(s.get("Category", "")).upper() and "IPO" in str(s.get("Category", "")).upper():
                normalized.extend([{**s, "Category": "M&A"}, {**s, "Category": "IPO"}])
            else: normalized.append(s)
                
        for s in normalized:
            r = t_lo.add_row().cells
            add_formatted_text(r[0].paragraphs[0].clear(), str(s.get("Category", "")))
            add_formatted_text(r[1].paragraphs[0].clear(), str(s.get("Probability", "")))
            add_formatted_text(r[2].paragraphs[0].clear(), str(s.get("Comment", "")))

    doc.add_heading("2-5. 주요 성장 지표", level=2)
    stats = mg.get("Export_and_Contract_Stats") or {}
    is_unit_str = str(data.get("Financial_Status", {}).get("Income_Statement", {}).get("Unit") or "백만원")
    fb_unit = "억원" if "억" in is_unit_str else "백만원"

    def _add_chart(doc, key, title, color):
        data_list = stats.get(key)
        if data_list:
            doc.add_paragraph(f"■ {title}")
            buf = create_basic_bar_chart_stream(data_list, title, color, fb_unit)
            if buf: doc.add_picture(buf, width=Inches(5.5))
        else:
            doc.add_paragraph(f"■ {title} : IR 자료 내 관련 데이터가 명시되어 있지 않습니다.")

    _add_chart(doc, "Export_Graph_Data", "수출 추이", COLOR_RED)
    _add_chart(doc, "Contract_Count_Graph_Data", "계약 건수", COLOR_YELLOW)
    _add_chart(doc, "Sales_Graph_Data", "매출 성장", COLOR_BLUE)
    doc.add_paragraph("\n")

def _render_tech_and_personnel(doc, data):
    doc.add_heading("3. 기술 경쟁력 및 파이프라인", level=1)
    tp = data.get("Technology_and_Pipeline") or {}
    
    doc.add_heading("3-1. Market Pain Points", level=2)
    for p_text in tp.get("Market_Pain_Points") or []: add_formatted_text(doc.add_paragraph(), f"• {p_text}")
    
    doc.add_heading("3-2. Solution & Core Tech", level=2)
    sol = tp.get("Solution_and_Core_Tech") or {}
    add_formatted_text(doc.add_paragraph(), f"핵심기술: {sol.get('Technology_Name')}")
    for k in sol.get("Key_Features") or []: add_formatted_text(doc.add_paragraph(), f"- {k}")

    doc.add_heading("3-3. 주요 파이프라인 개발 현황", level=2)
    pipe = tp.get("Pipeline_Development_Status") or {}
    add_formatted_text(doc.add_paragraph(), f"• 플랫폼 상세: {pipe.get('Core_Platform_Details')}")
    add_formatted_text(doc.add_paragraph(), f"• 위험도 분석: {pipe.get('Technical_Risk_Analysis')}")
    add_formatted_text(doc.add_paragraph(), f"• 결론: {pipe.get('Technical_Conclusion')}")

    add_smart_heading(doc, "4. 주요 인력 및 조직", level=1)
    kp = data.get("Key_Personnel") or {}
    
    add_smart_heading(doc, "4-1. 대표이사 레퍼런스", level=2)
    ceo = kp.get("CEO_Reference") or {}
    doc.add_paragraph(f"■ 성명: {ceo.get('Name', '')}").paragraph_format.keep_with_next = True
    doc.add_paragraph(f"■ 학력 및 경력:\n{ceo.get('Background_and_Education', '')}")
    doc.add_paragraph(f"■ 핵심 역량:\n{ceo.get('Core_Competency', '')}")
    doc.add_paragraph(f"■ 경영 철학:\n{ceo.get('Management_Philosophy', '')}")
    doc.add_paragraph(f"■ VC 관점 평가:\n{ceo.get('VC_Perspective_Evaluation', '')}")

    doc.add_heading("4-2. 조직 역량", level=2)
    team = kp.get("Team_Capability") or {}
    doc.add_paragraph("■ 핵심 임원진:")
    exec_list = team.get("Key_Executives")
    if isinstance(exec_list, str): doc.add_paragraph(f"- {exec_list}")
    elif isinstance(exec_list, list): 
        for ex in exec_list: doc.add_paragraph(f"- {ex}")
    else: doc.add_paragraph("- 정보 없음")
    doc.add_paragraph(f"■ 조직 강점:\n{team.get('Organization_Strengths', '')}")
    doc.add_paragraph(f"■ 자문단:\n{team.get('Advisory_Board', '')}")

    doc.add_heading("4-3. 주주명부", level=2)
    cap_table = kp.get("Cap_Table") or []
    if cap_table:
        t_cap = doc.add_table(rows=1, cols=3)
        t_cap.style = "Table Grid"
        for i, h in enumerate(["주주명", "보유 주식수", "지분율"]):
            cell = t_cap.rows[0].cells[i]
            cell.text = h
            set_cell_background(cell, "D9E1F2")
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs: run.font.bold, run.font.size = True, Pt(10)
        for item in cap_table:
            row = t_cap.add_row().cells
            row[0].text, row[1].text, row[2].text = str(item.get("Shareholder", "-")), str(item.get("Shares", "-")), str(item.get("Ratio", "-"))
            for cell in row:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs: run.font.size = Pt(10)
    else:
        doc.add_paragraph("IR 자료 내에 주주명부(지분 구조) 데이터가 명시되어 있지 않습니다.")
    doc.add_paragraph("")

def _render_risks_and_valuation(doc, data):
    add_smart_heading(doc, "5. 리스크 및 종합 투자 판단", level=1)
    add_smart_heading(doc, "5-1. 주요 리스크 및 대응", level=2)
    
    risks = data.get("Investment_Risks") or []
    if risks:
        for idx, r in enumerate(risks, 1):
            doc.add_paragraph().add_run(f"{idx}) {r.get('Risk_Title', '리스크 제목')}").bold = True
            doc.add_paragraph(f"  - {r.get('Risk_Detail', '')}")
            doc.add_paragraph(f"  --> {r.get('Mitigation_and_Outlook', '')}\n")
    else:
        doc.add_paragraph("- 주요 리스크 정보가 없습니다.")
            
    val_judge = data.get("Valuation_and_Judgment") or {}
    doc.add_heading("5-2. 밸류에이션 추정", level=2)
    
    val_table = val_judge.get("Valuation_Table") or []
    if val_table:
        vt = doc.add_table(rows=1, cols=4)
        vt.style = "Table Grid"
        for i, h in enumerate(["Round", "Pre-Money", "Post-Money", "Comment"]): vt.rows[0].cells[i].text = h
        apply_table_style(vt)
        for v in val_table:
            r = vt.add_row().cells
            r[0].text, r[1].text, r[2].text, r[3].text = str(v.get("Round", "-")), str(v.get("Pre_Money", "-")), str(v.get("Post_Money", "-")), str(v.get("Comment", "-"))
    
    logic = val_judge.get("Valuation_Logic_Detail") or {}
    p_logic = doc.add_paragraph()
    p_logic.add_run("■ 밸류에이션 산정 로직 (Data-Driven Analysis)\n").bold = True
    
    def format_list_with_count(raw_list, limit=20):
        if not raw_list: return "0개사 (해당 없음)"
        if not isinstance(raw_list, list): return str(raw_list)
        
        display_items = ", ".join([str(item) for item in raw_list[:limit]])
        return f"{len(raw_list)}개사 ({display_items}, ...)" if len(raw_list) > limit else f"{len(raw_list)}개사 ({display_items})"

    inds = logic.get("Step1_Industries") or logic.get("target_industries")
    p_logic.add_run(f"• 타겟 산업: {', '.join(inds) if isinstance(inds, list) else inds or '-'}\n\n")
    p_logic.add_run(f"• 1차 후보군: {format_list_with_count(logic.get('stage1_raw') or logic.get('Step2_Raw_Pool') or logic.get('Step2_Raw_List'))}\n")
    
    p_logic.add_run(f"• 2차 필터링(재무):\n")
    p_logic.add_run(f"  - 12월 결산월 기업: {format_list_with_count(logic.get('stage2_dec_passed') or logic.get('Step3_Dec_Filtered'))}\n")
    p_logic.add_run(f"  - 2개년도 전 흑자 기업: {format_list_with_count(logic.get('stage2_profit_passed') or logic.get('Step3_Profit_Filtered'))}\n\n")
    
    business_list = logic.get("stage3_business_passed") or logic.get("Step4_Business_Filtered")
    if business_list:
        p_logic.add_run(f"• 3단계 필터링(사업 유사성):\n  - 사업 유사성 통과: {format_list_with_count(business_list)}\n\n")
    
    details_info = logic.get("details") or {}
    similarity_scores = details_info.get("stage3_similarity") or logic.get("Stage3_Similarity_Scores")
    if similarity_scores:
        p_logic.add_run(f"■ 사업 유사도 상세 분석 (주요제품 구성 포함):\n(통과 기준: 유사도 0.30 이상)\n\n")
        for item in similarity_scores[:15]:
            status = "✅ 통과" if item.get("company") in (business_list or []) else "❌ 미통과"
            p_logic.add_run(f"{status} [{item.get('company', '-')}] 유사도 {item.get('score', 0.0):.2f}\n")
            p_logic.add_run(f"  주요제품: {item.get('main_products', '정보 없음')}\n  판단근거: {item.get('reason', '')}\n\n")
    
    requirements_check = details_info.get("stage4_requirements") or logic.get("Stage4_Requirements_Check")
    outlier_drop_count = 0
    if requirements_check:
        fetch_date = "N/A"
        if len(requirements_check) > 0:
            first = requirements_check[0]
            if isinstance(first, dict): fetch_date = first.get("info", {}).get("fetch_date", "N/A")
            elif len(first) >= 4 and isinstance(first[3], dict): fetch_date = first[3].get("fetch_date", "N/A")
        
        p_logic.add_run(f"■ 일반 요건 검증 결과 ({fetch_date} 기준):\n")
        for item in requirements_check[:30]:
            if isinstance(item, dict): comp, passed, reason, info = item.get("company"), item.get("passed"), item.get("reason", ""), item.get("info", {})
            elif isinstance(item, (tuple, list)) and len(item) >= 4: comp, passed, reason, info = item[0], item[1], item[2], item[3]
            else: continue
            
            if "아웃라이어" in reason or "MAX/MIN" in reason: outlier_drop_count += 1
            
            cap_val = info.get("market_cap", 0) if isinstance(info, dict) else 0
            cap_str = f"시총 {int(cap_val//10000)}조 {int(cap_val%10000):,}억" if cap_val >= 10000 and int(cap_val%10000) > 0 else (f"시총 {int(cap_val//10000)}조" if cap_val >= 10000 else f"시총 {int(cap_val):,}억") if cap_val > 0 else "시총 N/A"
            
            metrics = f" [PER {info.get('per', 'N/A')}, PBR {info.get('pbr', 'N/A')}, {cap_str}, EV/EBITDA {info.get('ev_ebitda', 'N/A')}]"
            p_logic.add_run(f"  - {comp}: {'✓ 통과' if passed else '✗ 탈락'}{f' ({reason})' if reason else ''}{metrics}\n")

    if len(logic.get("Step5_Final_Peers", [])) > 0:
        if outlier_drop_count == 0:
            p_logic.add_run("\n※ [비고] 1차 필터링을 통과한 Peer 기업 수가 한정적입니다, 통계적 유의성 확보(최소 표본 사수)를 위해 PER, EV/EBITDA 등의 MAX/MIN 아웃라이어 제거를 수행하지 않고 전원 최종 비교 기업에 편입하였습니다.\n").bold = True
        elif 0 < outlier_drop_count < 3:
            p_logic.add_run("\n※ [비고] 표본 수 확보를 위해 EV/EBITDA 지표의 아웃라이어 제거는 보류하였으며, 가치 산정에 직접적인 영향을 미치는 PER 지표의 MAX/MIN 기업만 제한적으로 제외하였습니다.\n").bold = True

    final_peers = logic.get("stage4_final_peers") or logic.get("Step5_Final_Peers") or logic.get("Step4_Final_Peers") or []
    p_logic.add_run(f"\n■ 최종 비교 기업 선정: {len(final_peers)}개사\n").bold = True
    for peer in final_peers: p_logic.add_run(f"  - {peer}\n") if final_peers else p_logic.add_run("  - 선정된 기업 없음\n")
    
    p_logic.add_run(f"\n■ Target 기업 가치 산출\n").bold = True
    p_logic.add_run(f"• 적용 Multiple: {logic.get('Applied_Multiple', '-')}\n")
    clean_rationale = re.sub(r'(원|배|주)([1-5]\.)', r'\1\n\2', logic.get('Calculation_Rationale', '-'))
    p_logic.add_run(f"• 산출 근거 및 수식:\n{clean_rationale}\n")

    scenarios = logic.get("Scenario_Valuation")
    if scenarios and isinstance(scenarios, list) and len(scenarios) > 0:
        p_note = doc.add_paragraph()
        run_note1 = p_note.add_run("※ 해당 시나리오는 낙관적(IR 추정치 100% 달성), 중립적(IR 추정치 70% 달성), 보수적(IR 추정치 50% 달성)을 기준으로 산정했습니다.\n")
        
        # 🚨 [핵심 수정] 40% 하드코딩을 지우고, valuation_agent에서 넘어온 변수를 연동합니다.
        d_pv = logic.get("Discount_Rate_PV", "50%")
        d_ipo = logic.get("Discount_Rate_IPO", "40%")
        
        run_note2 = p_note.add_run(f"※ 최종 예상공모가는 주당 평가가액 기준 공모 할인율({d_ipo})이 적용되었으며, 기업가치 현가 산출에는 라운드별 타겟 할인율({d_pv})이 반영되었습니다.")
        
        for r_note in (run_note1, run_note2): r_note.font.color.rgb, r_note.font.size = RGBColor(128, 128, 128), Pt(9)

        t_scen = doc.add_table(rows=1, cols=6)
        t_scen.style = "Table Grid"
        for i, h in enumerate(["시나리오", "실적 달성 가정", "추정 순이익 현가", "Target 기업가치", "주당 평가가액", "최종 예상공모가"]):
            cell = t_scen.rows[0].cells[i]
            cell.text = h
            set_cell_background(cell, "D9E1F2")
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs: run.font.bold, run.font.size = True, Pt(9)
        
        for scen in scenarios:
            row = t_scen.add_row().cells
            row[0].text, row[1].text, row[2].text, row[3].text, row[4].text, row[5].text= str(scen.get("Scenario", "")), str(scen.get("Ratio", "")), str(scen.get("PV", "")), str(scen.get("EV", "")), str(scen.get("Price", "")), str(scen.get("Final_Price", ""))
            for i in range(6):
                for p in row[i].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(9)
                        if i == 0: run.font.bold = True
        doc.add_paragraph()

    doc.add_heading("5-3. 종합 투자 판단", level=2)
    axes = val_judge.get("Three_Axis_Assessment") or {}
    add_formatted_text(doc.add_paragraph(), f"• 기술성: {axes.get('Technology_Rating') or '-'}")
    add_formatted_text(doc.add_paragraph(), f"• 성장성: {axes.get('Growth_Rating') or '-'}")
    add_formatted_text(doc.add_paragraph(), f"• 회수성: {axes.get('Exit_Rating') or '-'}")
    doc.add_paragraph(f"• 적합 투자자: {val_judge.get('Suitable_Investor_Type') or '-'}")

    doc.add_heading("5-4. 종합 결론", level=2)
    add_formatted_text(doc.add_paragraph(), data.get("Final_Conclusion") or "-")


# =========================================================
# 4. 메인 진입 함수 (Main Controller)
# =========================================================
def save_as_word_report(data, file_name, target_dir, original_file_name=""):
    if not data: return None
    if not os.path.exists(target_dir): os.makedirs(target_dir)
    
    doc = Document()
    header = data.get("Report_Header") or {}
    fin_data = data.get("Financial_Status", {})
    
    # [Cover]
    doc.add_paragraph("\n\n\n")
    t = doc.add_heading(f"{header.get('Company_Name', file_name)}\n투자 검토 보고서", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n" * 5)
    
    p = doc.add_paragraph(f"Date: {datetime.datetime.now().strftime('%Y. %m. %d')}\nAnalyst: {header.get('Analyst', 'LUCEN Investment Intelligence')}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_page_break()
    
    # [각 섹션 렌더링] - 분할된 모듈 순차 호출
    _render_executive_summary(doc, data, header, fin_data)
    _render_financial_and_investment(doc, data, file_name, original_file_name, fin_data)
    _render_market_and_growth(doc, data)
    _render_tech_and_personnel(doc, data)
    _render_risks_and_valuation(doc, data)

    # [최종 저장]
    out_path = os.path.join(target_dir, f"{file_name}_검토보고서.docx")
    doc.save(out_path)
    return out_path