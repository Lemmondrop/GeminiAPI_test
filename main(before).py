import os
import glob
import json
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from parser import extract_text_from_pdf
from processor import refine_to_json

# 한글 폰트 설정 (Windows 기준)
plt.rcParams['font.family'] = 'Malgun Gothic'
plt.rcParams['axes.unicode_minus'] = False

def create_financial_chart(financial_table, file_name, target_dir):
    """
    JSON의 Financial_Table 데이터를 바탕으로 막대 그래프 생성
    """
    try:
        if not financial_table or len(financial_table) < 2:
            return None

        # 데이터 파싱 (첫 줄은 연도, 두 번째 줄은 매출로 가정)
        years = financial_table[0][1:] # ['2023(A)', '2024(E)', ...]
        revenue_row = next((row for row in financial_table if "매출" in row[0]), None)

        if not revenue_row:
            return None

        # 콤마 제거 및 수치화
        values = []
        for v in revenue_row[1:]:
            clean_v = str(v).replace(',', '').replace('(', '-').replace(')', '').strip()
            values.append(float(clean_v) if clean_v else 0)

        # 그래프 그리기
        plt.figure(figsize=(8, 5))
        bars = plt.bar(years, values, color='#2E5A88', width=0.6) # 신뢰감 주는 네이비 톤

        # 수치 표시
        for bar in bars:
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                     f'{int(height):,}', ha='center', va='bottom', fontweight='bold')

        plt.title(f'연도별 매출 추이 및 전망 (단위: 억 원)', fontsize=14, pad=20)
        plt.grid(axis='y', linestyle='--', alpha=0.7)

        # 이미지 저장
        chart_path = os.path.join(target_dir, f"{file_name}_chart.png")
        plt.tight_layout()
        plt.savefig(chart_path)
        plt.close()
        return chart_path
    except Exception as e:
        print(f"   - [차트 생성 실패]: {str(e)}")
        return None

def save_as_word_report(data, file_name, target_dir):
    if not os.path.exists(target_dir):
        os.makedirs(target_dir)

    doc = Document()

    # 1. 문서 제목 (Title)
    header_info = data.get("Report_Header", {})
    company_name = header_info.get("Company_Name", file_name)
    title = doc.add_heading(f"투자 검토 보고서: {company_name}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2. 요약 정보 표 (에러 방지를 위해 기본 스타일 사용)
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Table Grid'

    rows_data = [
        ["기업명", header_info.get("Company_Name", "N/A")],
        ["담당 심사역", header_info.get("Analyst", "수석 심사역 유창훈")],
        ["투자 등급", header_info.get("Investment_Rating", "N/A")]
    ]

    for i, row in enumerate(rows_data):
        info_table.rows[i].cells[0].text = row[0]
        info_table.rows[i].cells[1].text = str(row[1])
    doc.add_paragraph() 

    # 3. 핵심 투자 하이라이트
    doc.add_heading('1. 핵심 투자 하이라이트 (Investment Thesis)', level=1)
    doc.add_paragraph(str(data.get('Investment_Thesis_Summary', '내용 없음')))

    # 4. 시장 배경 및 해결 방안 (복합 구조 방어 로직)
    doc.add_heading('2. 시장 배경 및 해결 방안', level=1)
    prob_sol = data.get('Problem_and_Solution', {})

    # Market Pain Points 처리
    doc.add_heading('▶ 시장의 주요 Pain Points', level=2)
    pain_points = prob_sol.get('Market_Pain_Points', [])
    if isinstance(pain_points, list):
        for point in pain_points:
            doc.add_paragraph(str(point), style='List Bullet')
    else:
        doc.add_paragraph(str(pain_points))

    # Solution Value Prop 처리
    doc.add_heading('▶ 동사의 해결책 및 가치 제안', level=2)
    solutions = prob_sol.get('Solution_Value_Prop', [])
    if isinstance(solutions, list):
        for sol in solutions:
            doc.add_paragraph(str(sol), style='List Bullet')

    # 5. 기술력 및 시장 규모
    doc.add_heading('3. 기술적 차별성 및 시장 규모', level=1)
    tech = data.get('Technology_and_Moat', {})

    # 에러 원인이었던 style='Strong'을 제거하고 직접 굵게 처리
    p = doc.add_paragraph()
    p.add_run(f"핵심 기술: {tech.get('Core_Technology_Name', 'N/A')}").bold = True

    tech_details = tech.get('Technical_Details', [])
    if isinstance(tech_details, list):
        for detail in tech_details:
            doc.add_paragraph(str(detail), style='List Bullet')

    doc.add_heading('▶ 시장 규모 (TAM/SAM/SOM)', level=2)
    market_opp = data.get('Market_Opportunity')
    if isinstance(market_opp, dict):
        doc.add_paragraph(str(market_opp.get('TAM_SAM_SOM', 'N/A')))
    else:
        doc.add_paragraph(str(market_opp))

    # 6. 주요 리스크 및 대응 전략 (표 형식)
    doc.add_heading('4. 주요 리스크 및 대응 전략', level=1)
    risks_data = data.get('Key_Risks_and_Mitigation')

    # 리스크가 리스트 형태인 경우 표로 생성
    if isinstance(risks_data, list) and len(risks_data) > 0:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '리스크 요인'
        hdr_cells[1].text = '대응 및 완화 전략'

        for r_item in risks_data:
            if isinstance(r_item, dict):
                row_cells = table.add_row().cells
                row_cells[0].text = str(r_item.get('Risk_Factor', r_item.get('risk', 'N/A')))
                row_cells[1].text = str(r_item.get('Mitigation_Strategy', r_item.get('mitigation', 'N/A')))
    else:
        doc.add_paragraph(str(risks_data))

    # 7. 최종 종합 의견
    doc.add_heading('5. 심사역 종합 의견', level=1)
    doc.add_paragraph(str(data.get('Final_Conclusion', data.get('Overall_Opinion', '내용 없음'))))

    # 파일 저장
    save_name = f"{file_name}_투자검토보고서.docx"
    report_path = os.path.join(target_dir, save_name)
    doc.save(report_path)
    return report_path

def main():
    input_dir = "data"
    output_dir = "output"
    report_dir = "output_report"

    for d in [output_dir, report_dir]:
        if not os.path.exists(d): os.makedirs(d)

    pdf_files = glob.glob(os.path.join(input_dir, "*.pdf"))
    if not pdf_files:
        print(f"[{input_dir}] 폴더에 PDF 파일이 없습니다.")
        return

    print(f"총 {len(pdf_files)}개의 파일을 처리를 시작합니다.")

    for file_path in pdf_files:
        file_name = os.path.splitext(os.path.basename(file_path))[0]
        json_path = os.path.join(output_dir, f"{file_name}_refined.json")

        print(f"\n>>> [처리 중] {file_name}")

        try:
            # --- 성공한 파일 건너뛰기 로직 추가 ---
            is_already_done = False
            if os.path.exists(json_path):
                with open(json_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                
                # JSON 내부에 에러가 없고 정상 데이터인 경우
                if isinstance(data, dict) and "error" not in data:
                    print(f"   - [Skip] 이미 분석된 정상 JSON이 존재합니다. Word 보고서만 다시 생성합니다.")
                    save_as_word_report(data, file_name, report_dir)
                    print(f"   - [성공] 보고서 저장 완료")
                    is_already_done = True
            
            if is_already_done:
                continue
            # ------------------------------------

            # 1. 텍스트 추출
            raw_text = extract_text_from_pdf(file_path)
            
            # 2. Gemini AI 분석 (투심보고서용 JSON 생성)
            # API 호출 직전 짧은 대기 (Rate Limit 방지)
            import time
            time.sleep(1) 
            refined_data = refine_to_json(raw_text[:30000])
            
            # 3. JSON 저장 (에러가 나더라도 결과 기록)
            with open(json_path, "w", encoding="utf-8") as f:
                json.dump(refined_data, f, ensure_ascii=False, indent=2)

            # 4. 워드 보고서 생성 로직
            if isinstance(refined_data, dict) and "error" not in refined_data:
                print(f"   - 투자검토보고서(Word) 생성 중...")
                report_path = save_as_word_report(refined_data, file_name, report_dir)
                print(f"   - [성공] 보고서 저장 완료: {report_path}")
            else:
                err_msg = refined_data.get('message') if isinstance(refined_data, dict) else "Data Format Error"
                print(f"   - [경고] API 오류로 보고서 생성 건너뜀: {err_msg[:50]}...")

        except Exception as e:
            print(f"   - [{file_name}] 처리 중 예외 발생: {str(e)}")

if __name__ == "__main__":
    main()